from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(parents=True, exist_ok=True)

PRD_PATH = OUT_DIR / "标策AI_产品需求文档_PRD_V1.0.docx"
FRONTEND_PATH = OUT_DIR / "标策AI_前端设计文档_V1.0.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "19332B"
GREEN = "315C4D"
INK = "1D2B26"
MUTED = "6E7772"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E4EEE7"
LIGHT_AMBER = "F5E8D4"
LIGHT_RED = "F2DDD6"
LIGHT_GRAY = "F2F4F7"
BORDER = "D7DDE2"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT_WESTERN = "Calibri"
FONT_CJK = "Microsoft YaHei"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 western: str = FONT_WESTERN, cjk: str = FONT_CJK):
    run.font.name = western
    if run._element.get_or_add_rPr().rFonts is None:
        run._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), western)
    rfonts.set(qn("w:hAnsi"), western)
    rfonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_style_font(style, size: float, color: str = INK, bold: bool = False,
                   western: str = FONT_WESTERN, cjk: str = FONT_CJK):
    style.font.name = western
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), western)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), western)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk)
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold


def configure_document(doc: Document, preset: str, running_title: str, doc_code: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10 if preset == "standard_business_brief" else 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 16 if preset == "standard_business_brief" else 18, 8 if preset == "standard_business_brief" else 10),
        "Heading 2": (13, BLUE, 12 if preset == "standard_business_brief" else 14, 6 if preset == "standard_business_brief" else 7),
        "Heading 3": (12, DARK_BLUE, 8 if preset == "standard_business_brief" else 10, 4 if preset == "standard_business_brief" else 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.05

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, 11, INK)
        style.paragraph_format.left_indent = Inches(0.5 if preset == "standard_business_brief" else 0.375)
        style.paragraph_format.first_line_indent = Inches(-0.25 if preset == "standard_business_brief" else -0.188)
        style.paragraph_format.space_after = Pt(8 if preset == "standard_business_brief" else 4)
        style.paragraph_format.line_spacing = 1.167 if preset == "standard_business_brief" else 1.25

    for style_name, size, color, bold in [
        ("Doc Kicker", 9, BLUE, True),
        ("Doc Title", 25, INK, True),
        ("Doc Subtitle", 13, MUTED, False),
        ("Lead", 12, INK, True),
        ("Small Note", 9, MUTED, False),
        ("Table Text", 9.2, INK, False),
        ("Table Header", 9.2, INK, True),
        ("Code Block", 9, DARK_BLUE, False),
    ]:
        if style_name not in [s.name for s in styles]:
            style = styles.add_style(style_name, 1)
        else:
            style = styles[style_name]
        set_style_font(style, size, color, bold, "Consolas" if style_name == "Code Block" else FONT_WESTERN)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.1

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run(f"{running_title}  |  {doc_code}")
    set_run_font(run, 8.5, MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    set_run_font(fp.add_run("内部评审稿 · 2026-08-11"), 8, MUTED)
    fp.add_run("\t")
    set_run_font(fp.add_run("第 "), 8, MUTED)
    add_field(fp, "PAGE")
    set_run_font(fp.add_run(" 页"), 8, MUTED)

    props = doc.core_properties
    props.title = running_title
    props.subject = "标策 AI 多智能体投标决策产品文档"
    props.keywords = "投标, 评标, 多智能体, 招标文件解析, 报价优化"
    props.author = "标策 AI 项目组"
    props.comments = "基于 2026-08-11 前端 Demo 生成的内部评审文档"


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_sep, text, fld_char_end])


def add_rule(paragraph, color: str = BORDER, size: int = 8):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_masthead(doc: Document, kicker: str, title: str, subtitle: str,
                 metadata: Sequence[tuple[str, str]]):
    p = doc.add_paragraph(style="Doc Kicker")
    p.paragraph_format.space_after = Pt(8)
    p.add_run(kicker.upper())

    p = doc.add_paragraph(style="Doc Title")
    p.paragraph_format.space_after = Pt(5)
    p.add_run(title)

    p = doc.add_paragraph(style="Doc Subtitle")
    p.paragraph_format.space_after = Pt(16)
    p.add_run(subtitle)

    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(f"{label}："), 10, INK, bold=True)
        set_run_font(p.add_run(value), 10, INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(14)
    add_rule(rule, GREEN, 14)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color: str = BORDER, size: int = 4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths_dxa: Sequence[int], header_fill: str = LIGHT_BLUE,
              status_columns: Iterable[int] = ()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = 0
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in status_columns else WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(p.add_run(value), 9.2, INK, bold=True)

    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.08
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in status_columns else WD_ALIGN_PARAGRAPH.LEFT
            color = INK
            if idx in status_columns:
                if value in ("P0", "高", "阻断", "未实现"):
                    color = RED
                elif value in ("P1", "中", "部分实现", "风险"):
                    color = GOLD
                elif value in ("已实现", "通过", "低"):
                    color = GREEN
            set_run_font(p.add_run(str(value)), 9.1, color, bold=idx in status_columns)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_GREEN,
                accent: str = GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, fill, 0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(label.upper()), 8.5, accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    set_run_font(p2.add_run(text), 10.5, INK, bold=True)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: Sequence[str]):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, BORDER, 4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(line), 8.7, DARK_BLUE, western="Consolas", cjk=FONT_CJK)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc: Document, text: str, bold_lead: str | None = None,
             style: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), 11, INK, bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]), 11, INK)
    else:
        set_run_font(p.add_run(text), 11, INK)
    return p


def add_bullets(doc: Document, items: Sequence[str | tuple[str, str]], numbered: bool = False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        if isinstance(item, tuple):
            label, body = item
            set_run_font(p.add_run(label), 11, INK, bold=True)
            set_run_font(p.add_run(body), 11, INK)
        else:
            set_run_font(p.add_run(item), 11, INK)


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_prd() -> Document:
    doc = Document()
    configure_document(doc, "standard_business_brief", "标策 AI 产品需求文档", "BC-AI-PRD-001")
    add_masthead(
        doc,
        "Product Requirements Document",
        "标策 AI 产品需求文档（PRD）",
        "合规优先的招标文件解析、企业证据匹配与多智能体投标决策系统",
        [
            ("版本", "V1.0"),
            ("状态", "内部评审稿"),
            ("基线", "2026-08-11 已部署前端 Demo"),
            ("目标读者", "产品、投标、技术、财务、法务与管理层"),
        ],
    )
    add_callout(doc, "产品一句话", "把每个项目的招标文件转化为可执行规则，再把公司材料转化为可验证证据，最后在合规和利润约束下模拟竞争并给出可解释的投标策略。")
    add_heading(doc, "30 秒读懂", 2)
    add_bullets(doc, [
        ("唯一规则源：", "评标条件、权重和公式来自当前项目招标文件，不使用全国统一的“万能评分表”。"),
        ("先门槛后评分：", "资格和实质性要求先按通过/不通过审查，有效投标才进入详细评审。"),
        ("双向文档匹配：", "招标要求与公司资质、业绩、人员和技术材料逐条映射，输出缺口和原文位置。"),
        ("多智能体竞演：", "A 代表我方，B/C/D 模拟不同竞争策略，确定性规则裁判负责淘汰、计分和排序。"),
        ("不承诺中标：", "系统提高决策质量与可解释性，最终结论仍需项目负责人和评标委员会依法确认。"),
    ])
    add_page_break(doc)

    add_heading(doc, "1. 产品背景与问题定义", 1)
    add_heading(doc, "1.1 业务背景", 2)
    add_para(doc, "投标决策通常同时涉及招标文件阅读、资格核验、技术响应、商务证明、成本测算和竞争报价。现实中这些信息分散在大量 PDF、Word、Excel 和扫描件中，且每个项目的资格条件、实质性要求、评分因素、权重和公式均可能不同。")
    add_heading(doc, "1.2 核心痛点", 2)
    add_bullets(doc, [
        "人工阅读耗时，关键的 ★ 条款、废标条件和更正公告容易遗漏。",
        "公司资质、业绩、人员证书和技术方案分散，难以快速判断是否满足。",
        "报价优化常与合规审查、技术得分和利润底线脱节，出现“低价但不可投”或“高分但不盈利”。",
        "传统模型只给结论，无法说明规则来源、证据位置、缺口责任人和不确定性。",
        "不同项目评分机制不同，复用固定权重会产生合规和决策风险。",
    ])
    add_heading(doc, "1.3 产品机会", 2)
    add_para(doc, "通过文档智能体、企业证据库、确定性规则引擎和竞标策略智能体组合，可以把“阅读文件—判断能否投—准备响应—选择报价—形成报告”串成一条可追溯的决策链。")

    add_heading(doc, "2. 产品目标与非目标", 1)
    add_table(doc,
        ["类别", "内容", "衡量方式"],
        [
            ["目标", "在数分钟内形成项目规则初稿和资格/技术缺口清单", "抽取覆盖率、人工复核耗时"],
            ["目标", "在公司利润底线内比较多种报价策略", "可行报价区间、胜率与期望利润"],
            ["目标", "每个结论均可追溯到招标条款与公司证据", "规则/证据定位覆盖率"],
            ["目标", "支持投标、技术、商务、财务共同复核", "任务闭环率、决策周期"],
            ["非目标", "预测或保证真实中标结果", "页面与报告持续提示边界"],
            ["非目标", "替代法务、投标负责人或评标委员会判断", "保留人工确认节点"],
            ["非目标", "模型自行发明评分权重或资格加分项", "规则版本必须来自项目文件"],
        ],
        [1150, 5600, 2610],
        status_columns=(),
    )

    add_heading(doc, "3. 用户与关键场景", 1)
    add_table(doc,
        ["角色", "主要任务", "最关心的问题", "核心权限"],
        [
            ["投标经理", "建立项目、上传文件、协调响应", "能不能投、缺什么、何时报价", "项目全流程编辑"],
            ["资质/商务专员", "维护证照、业绩和授权材料", "证明是否有效、是否满足格式", "企业证据库维护"],
            ["技术负责人", "准备偏离表、技术方案和演示证据", "哪些条款必须响应、预计得分", "技术响应与评分复核"],
            ["财务/报价负责人", "成本测算、毛利约束和报价审批", "利润底价、异常低价和期望收益", "成本与报价情景"],
            ["管理层", "Bid/No-Bid 与最终报价决策", "风险、胜率、收益和建议", "查看与审批"],
            ["法务/合规", "复核规则、声明与风险", "是否误读条款、是否留痕", "规则纠正与审计"],
        ],
        [1500, 2600, 3100, 2160],
    )

    add_heading(doc, "4. 产品原则", 1)
    add_bullets(doc, [
        ("项目专属：", "每次建项都重新解析规则；模板只用于初始化，不替代正式招标文件。"),
        ("确定性裁判：", "硬门槛、限价、计分公式和排序由规则引擎执行，生成式智能体不能改写。"),
        ("证据优先：", "能力、证书或业绩只有在招标文件列为要求且具备规定证明材料时才算满足或得分。"),
        ("版本可追溯：", "招标更正、人工纠错和证据替换都生成版本并触发重新评估。"),
        ("人工在环：", "高风险条款、低置信度抽取、异常低价和最终决策必须有人确认。"),
        ("概率非承诺：", "胜出概率是模拟结果，必须展示假设、竞争者画像与敏感性。"),
    ])
    add_callout(doc, "不可违反的顺序", "招标文件解析 → Bid/No-Bid → 资格审查 → 符合性审查 → 异常低价审查 → 详细评审 → 报价优化 → 候选与报告。", LIGHT_AMBER, GOLD)

    add_heading(doc, "5. 端到端业务流程", 1)
    flow = [
        ("创建项目与上传文件。", "上传招标主文件、采购需求、评分办法、附件和更正公告。"),
        ("抽取并锁定规则。", "识别资格条件、实质性要求、评分项、权重、公式、最高限价和证明要求。"),
        ("接入公司材料。", "上传或复用营业执照、资质、证书、业绩、人员、技术和服务材料。"),
        ("执行证据匹配。", "逐条输出已满足、部分满足、缺失、置信度及双向原文定位。"),
        ("作出 Bid/No-Bid。", "先判断门槛、工期、资源、成本和利润可行性，再决定是否参与。"),
        ("运行多智能体竞演。", "模拟我方与不同竞争策略，规则裁判执行淘汰、计分、排名和概率映射。"),
        ("优化报价并出报告。", "按中标概率、期望利润或稳健收益目标搜索报价，生成审批与封标待办。"),
    ]
    add_bullets(doc, [(f"步骤 {idx + 1}：{title}", f" {body}") for idx, (title, body) in enumerate(flow)], numbered=True)

    add_heading(doc, "6. 功能需求", 1)
    add_table(doc,
        ["编号", "模块", "需求摘要", "优先级", "Demo 状态"],
        [
            ["FR-01", "项目与文件接入", "批量接收招标文件、附件和公司材料，展示类型、大小、版本与处理状态", "P0", "部分实现"],
            ["FR-02", "招标规则解析", "抽取资格、★ 条款、评分权重、公式、限价、工期及证明要求", "P0", "模拟实现"],
            ["FR-03", "企业证据库", "沉淀资质、业绩、人员与技术证据，支持有效期、归属与复用", "P0", "未实现"],
            ["FR-04", "双向证据匹配", "输出要求—证据矩阵、覆盖率、缺口、置信度与出处", "P0", "模拟实现"],
            ["FR-05", "Bid/No-Bid", "综合门槛、资源、成本、利润和竞争力给出可解释建议", "P0", "已实现"],
            ["FR-06", "评标规则引擎", "按项目顺序执行资格、符合性、低价审查、计分和排序", "P0", "已实现"],
            ["FR-07", "多智能体竞演", "A 为我方，B/C/D 采用低价、技术溢价和均衡策略", "P0", "已实现"],
            ["FR-08", "报价优化", "在利润与限价约束下优化胜率、期望利润或稳健收益", "P0", "已实现"],
            ["FR-09", "缺口协同", "将缺口分配给商务、技术、财务或法务并跟踪闭环", "P1", "未实现"],
            ["FR-10", "报告与审计", "导出规则、假设、结果、风险、人工确认与版本记录", "P0", "部分实现"],
        ],
        [820, 1540, 4380, 900, 1720],
        status_columns=(3, 4),
    )

    detailed_requirements = [
        ("6.1 FR-01 项目与文件接入", [
            "支持 PDF、Word、Excel、图片及压缩附件；保留原文件、哈希、来源、上传人和时间。",
            "区分招标方文件与企业材料，允许主文件、附件、更正公告建立引用关系。",
            "上传后执行病毒扫描、格式校验、OCR/版面识别与处理状态展示。",
            "验收：至少可稳定处理 500 页 PDF、扫描件和包含表格的评分办法。",
        ]),
        ("6.2 FR-02 招标规则解析", [
            "抽取规则必须结构化为：类型、文本、强制性、分值、公式、所需证据、来源页码和置信度。",
            "低置信度、冲突条款和更正前后差异进入人工复核队列。",
            "规则确认后锁定版本；后续修改必须重新运行证据匹配与报价模拟。",
        ]),
        ("6.3 FR-03/04 企业证据库与匹配", [
            "企业证据包含证据类型、主体、有效期、项目范围、原文位置、可复用性和敏感级别。",
            "匹配状态至少包含：已满足、部分满足、缺失、不适用、待人工确认。",
            "匹配应支持一条要求对应多份证据，以及一份证据支持多个要求。",
            "强制项未完全满足时，界面必须阻止“预计可通过”的绿色结论。",
        ]),
        ("6.4 FR-05/06 决策与规则引擎", [
            "资格和实质性要求使用通过制；内部利润底线不得混入评委评分。",
            "异常低价触发审查不等于自动废标，需展示说明材料与人工判断状态。",
            "最低评标价法与综合评分法使用不同的排序路径。",
            "每次运行记录规则版本、输入、公式、淘汰原因和输出。",
        ]),
        ("6.5 FR-07/08 多智能体与报价优化", [
            "我方智能体使用企业能力、证据成熟度、成本和报价；竞争者智能体使用可配置策略画像。",
            "规则裁判为独立确定性模块，任何投标智能体均不能修改规则。",
            "优化目标支持最大中标概率、最大期望利润、稳健收益；只搜索满足限价和最低毛利率的报价。",
            "输出推荐区间、当前点、敏感性和主要不确定性，而非单一神谕式价格。",
        ]),
    ]
    for title, bullets in detailed_requirements:
        add_heading(doc, title, 2)
        add_bullets(doc, bullets)

    add_heading(doc, "7. 多智能体与规则架构", 1)
    add_code_block(doc, [
        "招标文件 ──> 文档规则智能体 ──> 项目规则版本 ─┐",
        "公司材料 ──> 企业证据智能体 ──> 证据索引 ───┤",
        "                                               ├─> 确定性规则裁判 ─> 有效投标集合",
        "我方策略 A + 竞争策略 B/C/D ───────────────────┘                     │",
        "                                                                         └─> 报价优化 ─> 决策报告",
    ])
    add_table(doc,
        ["智能体/模块", "职责", "允许做什么", "禁止做什么"],
        [
            ["规则智能体", "抽取资格、条款、权重和公式", "提出结构化规则与置信度", "自行确定项目权重"],
            ["资格智能体", "识别准入和废标风险", "匹配证照、声明和授权", "把资格条件当作额外加分"],
            ["技术/评分智能体", "识别技术响应与评分证据", "预测得分区间和缺口", "无证据时给满分"],
            ["企业证据智能体", "索引公司材料", "定位证据、有效期和主体", "替代人工确认真实性"],
            ["投标策略 A/B/C/D", "生成不同报价和响应策略", "调整报价、成熟度和假设", "绕过硬门槛"],
            ["确定性规则裁判", "执行淘汰、计分和排序", "复现招标文件公式", "生成或改写规则"],
        ],
        [1700, 2300, 2700, 2660],
    )

    add_heading(doc, "8. 核心决策逻辑", 1)
    add_heading(doc, "8.1 硬约束", 2)
    add_code_block(doc, [
        "有效投标 = 资格通过 AND 符合性通过 AND 报价不超过最高限价",
        "内部可行报价 = 有效投标 AND 毛利率 >= 公司最低毛利率",
        "最低利润报价 = 预计成本 / (1 - 最低毛利率)",
    ])
    add_heading(doc, "8.2 综合评分示例", 2)
    add_code_block(doc, [
        "价格得分 = 最低有效评标价 / 本投标评标价 × 价格权重",
        "模拟总分 = 价格得分 + 技术得分 + 商务得分 + 服务得分",
        "注意：实际项目必须替换为招标文件载明的公式。",
    ])
    add_heading(doc, "8.3 优化目标", 2)
    add_bullets(doc, [
        "中标概率优先：在利润底线内提高模拟胜出概率。",
        "期望利润优先：中标概率 ×（报价 - 预计成本）。",
        "稳健收益优先：在期望利润基础上扣除异常低价和不确定性风险。",
    ])

    add_heading(doc, "9. 核心数据对象", 1)
    add_table(doc,
        ["对象", "关键字段", "用途"],
        [
            ["Project", "项目编号、采购类别、状态、负责人", "承载一次投标决策"],
            ["Document", "文件、类型、来源、哈希、版本、处理状态", "管理原始招标/企业材料"],
            ["RuleVersion", "方法、限价、权重、公式、确认人", "锁定项目规则基线"],
            ["Requirement", "类型、强制性、分值、证据要求、来源", "资格/符合性/评分的原子单元"],
            ["Evidence", "类型、主体、有效期、原文定位、敏感级", "企业可复用证明材料"],
            ["EvidenceMatch", "要求、证据、状态、置信度、复核人", "形成双向匹配矩阵"],
            ["Scenario", "成本、报价、利润、竞争强度、目标", "一次竞标模拟输入"],
            ["AgentResult", "有效性、分项得分、排名、概率、原因", "多智能体竞演输出"],
            ["DecisionReport", "结论、假设、风险、待办、审批记录", "管理层决策与审计"],
        ],
        [1600, 4800, 2960],
    )

    add_heading(doc, "10. Demo 与正式版边界", 1)
    add_table(doc,
        ["能力", "当前前端 Demo", "正式版要求"],
        [
            ["文件上传", "记录文件名、大小和类型，不读取正文", "加密对象存储、病毒扫描、权限和版本"],
            ["文档解析", "定时动画 + 项目模拟规则", "OCR、版面解析、表格识别、LLM 抽取和校验"],
            ["公司证据", "模拟材料列表和匹配结果", "企业证据库、有效期、主体、全文/向量索引"],
            ["评标规则", "三个项目模板和确定性前端函数", "规则 DSL、版本管理、公式测试和人工确认"],
            ["竞争模拟", "四个固定画像 + 概率映射", "历史数据校准、场景抽样、置信区间和模型评估"],
            ["协作审计", "本地状态与 JSON 导出", "账号、角色、任务、审批、日志和留存策略"],
        ],
        [1700, 3250, 4410],
    )
    add_callout(doc, "关键边界", "当前 Demo 可以用于确认产品流程、页面和决策逻辑，不可用于对真实招标文件作正式合规判断。", LIGHT_RED, RED)

    add_heading(doc, "11. 非功能需求", 1)
    add_table(doc,
        ["领域", "要求", "验收参考"],
        [
            ["安全", "传输与静态加密、租户隔离、最小权限、敏感文件水印", "通过安全评审与渗透测试"],
            ["隐私", "文件用途告知、保留期限、删除能力、模型训练隔离", "可追踪数据生命周期"],
            ["可靠性", "解析任务可重试、幂等、断点续传，规则引擎结果可复现", "关键任务成功率 ≥ 99.5%"],
            ["性能", "上传即时反馈，常规项目在可接受时间内形成规则初稿", "100 页 PDF 初稿目标 ≤ 5 分钟"],
            ["可解释性", "每条判断含来源、证据、置信度、公式和版本", "关键结论追溯覆盖率 100%"],
            ["可访问性", "键盘操作、可见焦点、屏幕阅读器标签、颜色非唯一编码", "核心流程满足 WCAG 2.1 AA"],
        ],
        [1500, 5000, 2860],
    )

    add_heading(doc, "12. 产品指标", 1)
    add_table(doc,
        ["层级", "指标", "建议口径"],
        [
            ["北极星", "从项目创建到形成可复核 Bid/No-Bid 建议的中位时长", "按项目、采购类别分层统计"],
            ["效率", "招标文件人工阅读时间下降比例", "试点前后同类项目对比"],
            ["质量", "强制条款召回率、规则纠错率", "以人工金标准集评估"],
            ["覆盖", "强制要求证据匹配覆盖率", "已满足/全部强制要求"],
            ["协作", "封标前缺口按时闭环率", "按任务截止时间统计"],
            ["决策", "报价建议被采纳率与采纳原因", "区分胜率/利润/稳健目标"],
        ],
        [1450, 4850, 3060],
    )

    add_heading(doc, "13. 路线图", 1)
    add_table(doc,
        ["阶段", "范围", "退出条件"],
        [
            ["V0.2 · 当前", "前端流程、模拟上传解析、三类项目模板、四方竞演和报告", "完成产品评审与交互确认"],
            ["V0.5 · MVP", "真实文件存储/解析、规则确认、企业证据库、基础账号权限", "5–10 个历史项目回放通过"],
            ["V0.8 · 试点", "任务协作、审计、规则 DSL、历史竞争画像与评估", "真实试点项目稳定运行"],
            ["V1.0 · 生产", "安全合规、监控、容灾、模型评测、企业集成与运维", "生产验收与制度落地"],
        ],
        [1600, 4700, 3060],
    )

    add_heading(doc, "14. 主要风险与对策", 1)
    add_table(doc,
        ["风险", "影响", "控制措施", "等级"],
        [
            ["招标条款抽取错误", "误判资格或评分", "原文定位、置信度、双模型/规则校验、人工锁定", "高"],
            ["扫描件和复杂表格识别失败", "规则缺失", "OCR 质量检测、表格专用解析、失败告警", "高"],
            ["公司证据过期或主体不符", "资格不通过", "有效期和主体校验、到期提醒、人工复核", "高"],
            ["概率被误解为保证", "错误决策与合规风险", "展示假设、区间、敏感性和免责声明", "高"],
            ["竞争数据不足", "概率校准偏差", "场景法、保守先验、历史回放和持续评估", "中"],
            ["敏感投标文件泄露", "重大安全风险", "加密、隔离、审计、最小权限和删除策略", "高"],
        ],
        [2500, 2100, 3860, 900],
        status_columns=(3,),
    )

    add_heading(doc, "15. MVP 验收标准", 1)
    add_bullets(doc, [
        "用户可创建项目并上传招标文件与企业材料，查看处理状态和版本。",
        "系统能从约定测试集抽取资格、实质性要求、评分因素、权重、公式和限价，并定位原文。",
        "用户可纠正规则并锁定版本；规则变更后自动重跑匹配和模拟。",
        "系统输出逐条证据匹配、缺口、置信度和人工确认状态。",
        "任何强制项未满足时，不得输出“预计可进入详细评审”。",
        "报价优化只在最高限价和最低毛利率约束内搜索，并展示目标与假设。",
        "报告包含规则版本、淘汰原因、分项得分、推荐报价、风险、待办和免责声明。",
        "关键操作具备角色权限和审计记录；敏感文件可按项目删除。",
    ])

    add_heading(doc, "附录 A：术语", 1)
    add_table(doc,
        ["术语", "定义"],
        [
            ["资格审查", "判断投标人主体、信用、资质和其他准入条件是否符合。通常为通过制。"],
            ["符合性审查", "判断报价、签章、有效期、★ 条款等是否实质响应。通常为通过制。"],
            ["详细评审", "只对有效投标按项目公式进行价格、技术、商务、服务等评审。"],
            ["证据匹配", "把招标要求与企业可验证材料建立可追溯关系。"],
            ["规则版本", "经解析和人工确认后锁定的一组项目评标规则。"],
            ["期望利润", "模拟中标概率 × 中标后的预计利润。"],
        ],
        [1900, 7460],
    )
    return doc


def build_frontend_design() -> Document:
    doc = Document()
    configure_document(doc, "compact_reference_guide", "标策 AI 前端设计文档", "BC-AI-FE-001")
    add_masthead(
        doc,
        "Frontend Product Design Specification",
        "标策 AI 前端设计文档",
        "页面信息架构、状态模型、组件规范、交互流程与正式版接口建议",
        [
            ("版本", "V1.0"),
            ("状态", "与 2026-08-11 前端 Demo 对齐"),
            ("目标", "帮助产品、设计和研发快速理解并继续实现"),
            ("线上 Demo", "https://bid-agent-lab-20260806.breezy-toad-2233.chatgpt.site"),
        ],
    )
    add_callout(doc, "设计主线", "第一屏解释“项目规则不是统一表”；第一个工作区完成文件上传与证据匹配；只有通过门槛后，用户才进入多智能体竞标与报价优化。")
    add_heading(doc, "快速索引", 2)
    add_bullets(doc, [
        "页面结构：Header + 项目 Hero + 七步流程 + 五个工作区 Tab。",
        "首要任务：上传招标方文件和公司材料，执行模拟解析并查看匹配矩阵。",
        "决策任务：设置成本、毛利、证据成熟度和竞争强度，运行四方竞演。",
        "解释任务：查看规则门槛、评分证据和最终决策报告。",
        "当前实现：单页面 React 客户端状态；正式版需拆分领域模块并接入后端任务。",
    ])
    add_page_break(doc)

    add_heading(doc, "1. 设计目标与原则", 1)
    add_table(doc,
        ["原则", "前端表达", "避免"],
        [
            ["合规优先", "门槛状态始终先于分数和概率出现", "只展示“最优报价”而隐藏废标风险"],
            ["项目专属", "项目模板切换后规则、权重、证据与公式同步变化", "固定通用权重"],
            ["可追溯", "规则和证据显示页码、章节、置信度与版本", "无来源的模型结论"],
            ["渐进披露", "先给结论与缺口，再展开公式、矩阵和解释", "第一屏堆满专业细节"],
            ["不确定性诚实", "概率标注“模拟”，异常低价标注“待审查”", "使用保证性语言"],
            ["协作导向", "缺口可以转为待办并回流模拟", "仅生成静态报告"],
        ],
        [1650, 4300, 3410],
    )

    add_heading(doc, "2. 信息架构", 1)
    add_code_block(doc, [
        "全局 Header：品牌 / 当前项目状态 / 方法边界 / 报告导出",
        "└─ 项目 Hero：价值主张 + 项目模板 + 文件处理状态 + 上传入口",
        "   └─ 七步流程：解析 → Bid/No-Bid → 资格 → 符合性 → 评分 → 报价 → 报告",
        "      ├─ Tab 1 文档上传与解析（默认）",
        "      ├─ Tab 2 决策驾驶舱",
        "      ├─ Tab 3 规则与门槛",
        "      ├─ Tab 4 评分证据",
        "      └─ Tab 5 决策报告",
    ])
    add_heading(doc, "2.1 导航规则", 2)
    add_bullets(doc, [
        "首次进入默认打开“文档上传与解析”，确保核心能力可被发现。",
        "Hero 中的“上传并解析招标文件”可回到文档工作区。",
        "解析结果中的“将结果带入竞标模拟”切换到决策驾驶舱，并注入门槛和成熟度状态。",
        "项目模板切换会重置文件、解析结果、竞标参数和规则版本，避免跨项目污染。",
        "“方法与边界”使用模态框，持续解释模型与法律边界。",
    ])

    add_heading(doc, "3. 页面全局结构", 1)
    add_table(doc,
        ["区域", "内容", "主要操作", "状态"],
        [
            ["Header", "品牌、规则引擎状态、项目编号", "方法说明、报告导出", "固定顶部"],
            ["Hero", "主张、原则、当前项目模板与预算", "切换模板、进入上传", "项目级"],
            ["Workflow", "七阶段业务流程", "当前版本只做进度表达", "全局"],
            ["Tabbar", "五个工作区及数量徽标", "切换工作区", "粘性心智导航"],
            ["Workspace", "上传、规则、竞演、证据或报告", "完成具体任务", "随 Tab 切换"],
            ["Footer", "Demo 边界与免责声明", "无", "全局"],
        ],
        [1450, 3500, 2500, 1910],
    )

    add_heading(doc, "4. 工作区 1：文档上传与解析", 1)
    add_heading(doc, "4.1 目标", 2)
    add_para(doc, "让用户明确区分“招标方规则源”和“本公司证据源”，完成双向上传、解析进度查看、规则确认和证据缺口判断。")
    add_heading(doc, "4.2 页面模块", 2)
    add_table(doc,
        ["模块", "说明", "关键交互", "验收重点"],
        [
            ["处理边界提示", "说明 Demo 不读取真实正文", "无", "不能误导用户已真实解析"],
            ["招标方文件区", "主文件、需求、评分办法、附件、更正", "点击/拖拽、批量、移除", "显示类型、大小、处理状态"],
            ["公司材料区", "资质、业绩、团队、技术、服务", "点击/拖拽、批量、移除", "与招标文件视觉区分"],
            ["解析流水线", "五类文档智能体及进度", "开始、重新解析、载入演示", "禁用态和阶段文案准确"],
            ["规则提取卡", "方法、限价、硬门槛、评分结构", "进入完整规则页", "显示来源与模拟置信度"],
            ["匹配摘要", "预计是否过门、四类覆盖率", "模拟补齐、带入竞演", "强制缺口必须为红色阻断"],
            ["匹配矩阵", "要求、证据、材料定位、状态、置信度", "横向滚动", "行级可追溯"],
        ],
        [1800, 3050, 2300, 2210],
    )
    add_heading(doc, "4.3 上传状态模型", 2)
    add_code_block(doc, [
        "EMPTY ──选择/拖拽──> READY ──开始解析──> PARSING ──成功──> DONE",
        "  ↑                       │                         │",
        "  └────移除全部文件────────┘                         ├─重新解析──> PARSING",
        "                                                    └─规则变更──> STALE",
        "正式版补充：UPLOADING / SCANNING / FAILED / CANCELED / STALE",
    ])
    add_heading(doc, "4.4 文件交互规范", 2)
    add_bullets(doc, [
        "接受 PDF、DOC/DOCX、XLS/XLSX、PNG/JPG；正式版需要展示单文件和项目总量限制。",
        "拖入时边框、底色和文案变化；离开或放下后恢复。",
        "同名同大小文件在当前 Demo 去重；正式版应以哈希去重并保留版本。",
        "解析中禁止重复提交；允许用户继续浏览其他工作区，但项目状态持续可见。",
        "失败必须显示可操作原因：格式不支持、密码保护、扫描质量差、解析超时或权限不足。",
    ])

    add_heading(doc, "5. 工作区 2：决策驾驶舱", 1)
    add_heading(doc, "5.1 左侧模拟输入", 2)
    add_table(doc,
        ["输入", "控件", "业务含义", "联动"],
        [
            ["我方投标总价", "滑杆 + 数字输入", "当前拟报价", "重算价格分、利润、排名和概率"],
            ["履约成本", "数字输入", "公司内部预计成本", "重算利润底价和期望利润"],
            ["最低毛利率", "滑杆", "公司硬约束", "限定优化搜索区间"],
            ["优化目标", "三段按钮", "胜率/期望利润/稳健收益", "改变推荐报价"],
            ["证据成熟度", "滑杆", "非价格项可验证程度", "影响我方非价格得分"],
            ["市场竞争强度", "滑杆", "竞争者报价压力", "扰动 B/C/D 报价"],
            ["门槛开关", "三组开关", "资格、★ 条款、低价说明", "决定有效性与 Bid/No-Bid"],
        ],
        [1600, 1700, 3250, 2810],
    )
    add_heading(doc, "5.2 结果层级", 2)
    add_bullets(doc, [
        ("第一层：", "GO / REVIEW / NO BID、建议报价、模拟胜率和预计毛利。"),
        ("第二层：", "A/B/C/D 四方卡片，展示报价、有效性、分项得分、排名和胜率。"),
        ("第三层：", "报价—胜率—期望利润前沿图，标记当前价和建议价。"),
        ("解释层：", "规则裁判执行顺序和每个智能体的淘汰/得分原因。"),
    ])
    add_callout(doc, "视觉优先级", "废标/不可行状态必须压过分数与胜率。无可行解时，建议报价、胜率和毛利显示空状态，不伪造数值。", LIGHT_RED, RED)

    add_heading(doc, "6. 工作区 3–5", 1)
    add_table(doc,
        ["工作区", "用户问题", "核心模块", "主要出口"],
        [
            ["规则与门槛", "这个项目到底怎么评？", "规则画像、资格/符合性清单、评分结构、公式、四方门槛矩阵", "回到证据或模拟"],
            ["评分证据", "我方为什么能拿这些分？", "证据成熟度、评分项—材料—页码—预测得分—缺口", "提高成熟度、创建补缺任务"],
            ["决策报告", "管理层应该如何决策？", "Bid/No-Bid、建议报价、依据、封标动作、风险边界", "复制摘要、导出 JSON/正式报告"],
        ],
        [1600, 2500, 3560, 1700],
    )

    add_heading(doc, "7. 关键交互流程", 1)
    add_heading(doc, "7.1 首次使用", 2)
    add_bullets(doc, [
        "选择或确认项目模板。",
        "上传招标方文件；如需判断我方资格，再上传公司材料。",
        "开始解析，查看规则提取结果与匹配缺口。",
        "人工确认规则；当前 Demo 可用“模拟补齐缺口”体验闭环。",
        "将解析结果带入竞标模拟，设置成本、毛利和目标。",
        "运行竞演，查看报价前沿并导出决策报告。",
    ], numbered=True)
    add_heading(doc, "7.2 缺口闭环", 2)
    add_code_block(doc, [
        "缺口发现 → 指定责任人/截止时间 → 上传补充证据 → 重新匹配",
        "        → 人工确认 → 更新资格/符合性状态 → 自动重跑报价模拟",
    ])
    add_heading(doc, "7.3 项目切换", 2)
    add_para(doc, "切换模板会清空当前文件和解析状态，并恢复该模板默认成本、毛利和报价。正式版若存在未保存内容，应先弹出确认并允许复制为新项目。")

    add_heading(doc, "8. 前端状态与数据映射", 1)
    add_table(doc,
        ["状态域", "当前变量/对象", "影响区域", "正式版来源"],
        [
            ["项目", "projectId / ProjectTemplate", "全局规则、模板和竞手", "Project API + RuleVersion"],
            ["文档", "tenderFiles / companyFiles", "上传区与数量徽标", "Document API + 对象存储"],
            ["解析", "parseState / progress / stage", "流水线与结果可见性", "ParseJob 轮询/推送"],
            ["证据匹配", "DocumentMatch[] / gapsClosed", "覆盖率、门槛和矩阵", "EvidenceMatch API"],
            ["我方约束", "ourBid / cost / minMargin", "利润、可行区间和优化", "Scenario 草稿"],
            ["模型假设", "readiness / marketPressure / objective", "得分、竞手和推荐价", "Scenario 参数"],
            ["评标结果", "Evaluation / AgentResult[]", "卡片、排名、图表和报告", "SimulationJob 结果"],
        ],
        [1550, 2850, 2450, 2510],
    )

    add_heading(doc, "9. 组件设计", 1)
    add_table(doc,
        ["组件", "职责", "关键 Props/事件", "复用位置"],
        [
            ["ProjectSelector", "切换项目规则上下文", "value, options, onChange", "Hero"],
            ["DocumentDropzone", "文件选择、拖拽、校验、列表", "kind, accept, files, onAdd, onRemove", "招标/公司上传"],
            ["AnalysisPipeline", "展示任务阶段和进度", "status, progress, stage, agents", "文档解析"],
            ["RuleSummaryCard", "展示方法、限价、门槛或评分", "label, value, source", "解析/规则页"],
            ["CoverageMetric", "覆盖率与进度条", "label, value, tone", "匹配摘要"],
            ["EvidenceMatchTable", "逐条要求—证据映射", "rows, filters, onReview", "解析/证据页"],
            ["ScenarioControls", "报价与内部约束输入", "scenario, onChange", "驾驶舱"],
            ["AgentCard", "单个投标智能体结果", "agent, method, rank", "竞演沙盘"],
            ["ScenarioChart", "胜率/收益前沿", "points, current, recommended", "驾驶舱"],
            ["DecisionReport", "结论、依据、风险与导出", "report, onExport", "报告页"],
        ],
        [1750, 2850, 2950, 1810],
    )

    add_heading(doc, "10. 视觉系统", 1)
    add_table(doc,
        ["Token", "当前值", "用途"],
        [
            ["Ink", "#1D2B26", "导航、主文字、深色决策条"],
            ["Canvas", "#F1EEE6", "全局暖灰背景"],
            ["Paper", "#FFFDF7", "卡片和文档感表面"],
            ["Green", "#315C4D", "通过、规则锁定、主曲线"],
            ["Blue", "#3568A8", "企业证据、期望利润、信息态"],
            ["Amber", "#D28C3C", "待审查、提醒、部分满足"],
            ["Rust/Red", "#C76B4F / #A8483E", "缺失、不通过、淘汰"],
            ["Radius", "6–12 px", "输入、卡片与状态标签"],
            ["Body font", "Noto Sans SC / 微软雅黑 / 系统字体", "中文界面正文"],
            ["Display font", "宋体回退", "主标题与决策数字"],
        ],
        [1900, 2600, 4860],
    )
    add_heading(doc, "10.1 状态颜色规则", 2)
    add_bullets(doc, [
        "绿色只表示已通过/已满足，不用于“概率较高”这类不确定结论。",
        "黄色表示需要审查、部分满足或不确定，不等于失败。",
        "红色表示阻断、缺失、不通过或超限；同时提供文字和图标，不只依赖颜色。",
        "蓝色表示信息、证据来源和辅助曲线，不参与通过/失败语义。",
    ])

    add_heading(doc, "11. 响应式与可访问性", 1)
    add_table(doc,
        ["断点", "布局策略", "注意事项"],
        [
            ["> 1240 px", "驾驶舱左侧控制栏 + 右侧结果；四张 Agent 卡并排", "保证图表和矩阵宽度"],
            ["960–1240 px", "控制栏转为横向双列；Agent 卡两列", "工作流和表格可横向滚动"],
            ["650–960 px", "上传区、规则卡和摘要改为单列/双列", "主要操作保持在模块顶部"],
            ["< 650 px", "全单列；按钮满宽；表格横向滚动", "触控目标至少约 44 px"],
        ],
        [1600, 4600, 3160],
    )
    add_bullets(doc, [
        "所有输入具备可见 label；文件移除按钮包含文件名 aria-label。",
        "拖拽上传必须同时提供键盘可用的文件选择入口。",
        "Tab、按钮、滑杆、开关和模态框具备可见焦点与合理的 Tab 顺序。",
        "Canvas 图表提供文本化 aria-label；正式版补充数据表或摘要作为等价替代。",
        "支持 prefers-reduced-motion；解析动画不得成为获取信息的唯一方式。",
    ])

    add_heading(doc, "12. 当前前端技术结构", 1)
    add_table(doc,
        ["文件/层", "当前职责", "建议演进"],
        [
            ["app/page.tsx", "模板数据、规则函数、状态、全部页面 JSX", "拆分 features/documents、rules、simulation、report"],
            ["app/globals.css", "完整视觉系统和响应式样式", "拆分 token、layout、component 样式或 CSS Modules"],
            ["app/layout.tsx", "中文元数据与社交预览", "根据请求 Host 生成绝对 OG URL"],
            ["前端内存状态", "文件元数据、解析动画、匹配和模拟", "服务端持久化 + Query 缓存 + URL/项目状态"],
            ["确定性函数", "evaluateScenario、优化搜索、覆盖率", "提取为共享规则包并增加单元/性质测试"],
            ["Canvas", "胜率和期望利润曲线", "封装图表组件并增加键盘/数据表替代"],
        ],
        [1900, 3400, 4060],
    )
    add_heading(doc, "12.1 推荐目录", 2)
    add_code_block(doc, [
        "app/",
        "  page.tsx                     # 页面装配与路由上下文",
        "features/documents/            # 上传、任务进度、规则提取、证据匹配",
        "features/rules/                # RuleVersion、门槛、公式与版本确认",
        "features/simulation/           # 场景输入、Agent 结果、优化和图表",
        "features/report/               # 决策报告、导出和审批",
        "components/ui/                 # Button、StatusTag、Card、Table、Modal",
        "lib/contracts/                 # API 类型和 Schema",
        "lib/rule-engine/               # 纯函数规则引擎（前后端共享）",
    ])

    add_heading(doc, "13. 正式版接口建议", 1)
    add_table(doc,
        ["方法", "路径", "用途", "前端关键状态"],
        [
            ["POST", "/projects", "创建投标项目", "creating / ready / failed"],
            ["POST", "/projects/{id}/documents", "初始化上传并返回签名地址", "uploading / scanning"],
            ["POST", "/projects/{id}/parse-jobs", "发起规则/证据解析", "queued / running / done / failed"],
            ["GET", "/parse-jobs/{id}", "查询阶段、进度和错误", "轮询或 SSE"],
            ["GET", "/projects/{id}/rule-versions/latest", "获取当前规则版本", "draft / confirmed / stale"],
            ["PATCH", "/requirements/{id}", "人工纠正规则并记录原因", "saving / saved / conflict"],
            ["GET", "/projects/{id}/evidence-matches", "获取匹配矩阵", "loading / ready / error"],
            ["POST", "/projects/{id}/scenarios", "运行竞标与报价优化", "queued / running / ready"],
            ["POST", "/projects/{id}/reports", "生成决策报告", "generating / downloadable"],
        ],
        [900, 3550, 2850, 2060],
    )
    add_heading(doc, "13.1 核心响应示例", 2)
    add_code_block(doc, [
        "Requirement {",
        "  id, kind, text, mandatory, score, formula,",
        "  requiredEvidence[], tenderSource { documentId, page, section },",
        "  confidence, reviewStatus, ruleVersionId",
        "}",
        "EvidenceMatch { requirementId, evidenceId, status, confidence, rationale, reviewer }",
        "ScenarioResult { validBids[], eliminatedBids[], ranking[], recommendation, assumptions[] }",
    ])

    add_heading(doc, "14. 错误与空状态", 1)
    add_table(doc,
        ["场景", "页面反馈", "恢复动作"],
        [
            ["未上传招标文件", "解析按钮禁用并提示最低前置条件", "上传文件或载入演示材料"],
            ["只上传招标文件", "可展示规则；证据矩阵全部标记未匹配", "上传公司材料"],
            ["文件不支持/加密", "文件行显示错误原因，不进入解析", "替换文件或输入密码后重传"],
            ["解析失败", "保留已上传文件和失败阶段", "重试、下载错误详情或人工录入"],
            ["规则有冲突", "黄色阻断，展示冲突来源", "人工选择有效版本并确认"],
            ["强制项缺失", "红色 GATE GAP，禁止绿色通过结论", "补充证据或作 NO BID"],
            ["利润底价高于限价", "NO BID，无推荐报价", "重估成本/毛利或退出项目"],
            ["网络中断", "保留本地草稿并提示同步状态", "自动重连或手动重试"],
        ],
        [2200, 4320, 2840],
    )

    add_heading(doc, "15. 埋点建议", 1)
    add_table(doc,
        ["事件", "触发", "关键属性"],
        [
            ["project_template_changed", "切换项目模板", "from, to"],
            ["document_added", "添加文件", "kind, extension, sizeBucket"],
            ["parse_started/completed/failed", "解析任务状态变化", "projectId, duration, stage, errorCode"],
            ["match_gap_reviewed", "用户查看或确认缺口", "requirementKind, status"],
            ["analysis_applied_to_simulation", "将匹配结果带入竞演", "coverage, hardGapCount"],
            ["scenario_parameter_changed", "调整报价/成本/目标", "field, previous, next"],
            ["simulation_run", "运行竞演", "objective, validCount, recommendation"],
            ["report_exported", "导出报告", "format, bidDecision"],
        ],
        [2900, 3000, 3460],
    )

    add_heading(doc, "16. 前端验收清单", 1)
    add_bullets(doc, [
        "首次进入能在首屏和默认 Tab 发现文档上传入口。",
        "两类文件区视觉和语义明确，点击、拖拽、批量添加和移除均可用。",
        "无招标文件时不能开始解析；解析中按钮状态、阶段和进度一致。",
        "完成后展示项目方法、限价、门槛、评分结构和报价公式。",
        "公司材料缺失时，匹配结果不得误报已满足；强制缺口使用阻断状态。",
        "“将结果带入竞标模拟”能同步资格、符合性和证据成熟度。",
        "切换项目模板后，文件、解析、规则和场景状态不串线。",
        "桌面、平板、手机均可完成核心流程；表格允许横向滚动。",
        "键盘焦点清晰、控件有标签、状态不只依赖颜色。",
        "所有概率、解析和法律判断均标注模拟/待人工复核边界。",
    ])

    add_heading(doc, "附录 A：当前 Demo 页面映射", 1)
    add_table(doc,
        ["页面元素", "实现位置", "备注"],
        [
            ["五个工作区与业务数据", "app/page.tsx", "单文件实现，便于 Demo，正式版建议拆分"],
            ["项目模板与匹配模拟数据", "app/page.tsx", "IT 服务、设备采购、智能化工程"],
            ["视觉与响应式", "app/globals.css", "暖纸张 + 深绿 + 蓝/黄/红状态"],
            ["站点元数据", "app/layout.tsx", "中文标题、描述和社交预览"],
            ["社交预览图", "public/bid-strategy-social.png", "四条策略路径汇入规则决策"],
        ],
        [2300, 2650, 4410],
    )
    return doc


def save_document(doc: Document, path: Path):
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
    doc.save(path)


def main():
    save_document(build_prd(), PRD_PATH)
    save_document(build_frontend_design(), FRONTEND_PATH)
    print(PRD_PATH)
    print(FRONTEND_PATH)


if __name__ == "__main__":
    main()
